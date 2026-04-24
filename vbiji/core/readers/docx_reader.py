"""Word 文档 (.docx) 读取器"""

from pathlib import Path
from docx import Document as DocxDocument
from vbiji.core.readers.base import BaseReader
from vbiji.core.types import Document as DocDocument


class DocxReader(BaseReader):
    """Word 文档读取器

    使用 python-docx 提取 .docx 文件中的段落文本
    """

    format = "docx"

    def supports(self, path: str) -> bool:
        """判断是否支持该文件

        Args:
            path: 文件路径

        Returns:
            bool: 如果是 .docx 文件返回 True
        """
        return Path(path).suffix.lower() == ".docx"

    def read(self, path: str) -> DocDocument:
        """读取 Word 文档

        Args:
            path: .docx 文件路径

        Returns:
            DocDocument: 包含文档文本内容的文档对象

        Raises:
            FileNotFoundError: 文件不存在时
        """
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        doc = DocxDocument(str(file_path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        content = "\n".join(paragraphs)
        metadata = {
            "paragraphs": len(paragraphs),
        }

        return DocDocument(
            content=content,
            title=file_path.stem,
            format=self.format,
            path=str(file_path.absolute()),
            metadata=metadata,
        )